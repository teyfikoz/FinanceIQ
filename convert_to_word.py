"""
Convert FinanceIQ Pro Markdown Documentation to Word Format
Creates a professionally formatted .docx file with proper styling
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_word_document():
    """Convert Markdown documentation to Word format"""

    # Create document
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Read Markdown file
    with open('FINANCEIQ_PRO_COMPREHENSIVE_DOCUMENTATION.md', 'r', encoding='utf-8') as f:
        content = f.read()

    # Split into lines
    lines = content.split('\n')

    # Process each line
    in_code_block = False
    in_table = False
    table_headers = []
    table_rows = []

    for line in lines:
        # Code blocks
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
            continue

        # Heading 1 (# )
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            heading = doc.add_heading(text, level=1)
            heading.paragraph_format.space_before = Pt(24)
            heading.paragraph_format.space_after = Pt(12)
            continue

        # Heading 2 (## )
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            heading = doc.add_heading(text, level=2)
            heading.paragraph_format.space_before = Pt(18)
            heading.paragraph_format.space_after = Pt(6)
            continue

        # Heading 3 (### )
        if line.startswith('### ') and not line.startswith('#### '):
            text = line[4:].strip()
            heading = doc.add_heading(text, level=3)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            continue

        # Heading 4 (#### )
        if line.startswith('#### '):
            text = line[5:].strip()
            heading = doc.add_heading(text, level=4)
            continue

        # Tables
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_headers = [cell.strip() for cell in line.split('|')[1:-1]]
                continue
            elif line.strip().replace('|', '').replace('-', '').strip() == '':
                # Skip separator row
                continue
            else:
                # Table data row
                row_data = [cell.strip() for cell in line.split('|')[1:-1]]
                table_rows.append(row_data)
                continue
        else:
            # End of table
            if in_table and table_headers and table_rows:
                # Create table
                table = doc.add_table(rows=len(table_rows) + 1, cols=len(table_headers))
                table.style = 'Light Grid Accent 1'

                # Header row
                header_cells = table.rows[0].cells
                for i, header in enumerate(table_headers):
                    header_cells[i].text = header
                    # Bold header
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # Data rows
                for row_idx, row_data in enumerate(table_rows):
                    cells = table.rows[row_idx + 1].cells
                    for col_idx, cell_data in enumerate(row_data):
                        if col_idx < len(cells):
                            cells[col_idx].text = cell_data

                # Reset table state
                in_table = False
                table_headers = []
                table_rows = []

                # Add spacing after table
                doc.add_paragraph()
                continue

        # Horizontal rule
        if line.strip() == '---':
            doc.add_paragraph('_' * 80)
            continue

        # Bullet lists (- or *)
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # Count indentation
            indent_level = (len(line) - len(line.lstrip())) // 2
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * indent_level)
            continue

        # Numbered lists
        if re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            indent_level = (len(line) - len(line.lstrip())) // 2
            p = doc.add_paragraph(text, style='List Number')
            p.paragraph_format.left_indent = Inches(0.25 * indent_level)
            continue

        # Bold text (**text**)
        if '**' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:  # Odd indices are bold
                    run.font.bold = True
            continue

        # Inline code (`code`)
        if '`' in line and not line.startswith('```'):
            p = doc.add_paragraph()
            parts = line.split('`')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:  # Odd indices are code
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
            continue

        # Empty lines
        if line.strip() == '':
            doc.add_paragraph()
            continue

        # Regular paragraph
        doc.add_paragraph(line)

    # Add cover page
    doc.paragraphs[0].insert_paragraph_before('')
    cover = doc.paragraphs[0]
    cover.text = 'FinanceIQ Pro'
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cover.runs:
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = RGBColor(102, 126, 234)

    cover.paragraph_format.space_after = Pt(12)

    # Add subtitle
    subtitle = doc.paragraphs[1].insert_paragraph_before('Bloomberg Terminal Seviyesinde Portföy Analitik Platformu')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(118, 75, 162)

    subtitle.paragraph_format.space_after = Pt(24)

    # Add date
    date_para = doc.paragraphs[2].insert_paragraph_before('Comprehensive Documentation')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date_para.runs:
        run.font.size = Pt(14)
        run.font.italic = True

    date_para.paragraph_format.space_after = Pt(6)

    # Add version
    version_para = doc.paragraphs[3].insert_paragraph_before('Version 1.4 | January 2025')
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in version_para.runs:
        run.font.size = Pt(12)

    version_para.paragraph_format.space_after = Pt(48)

    # Save document
    doc.save('FINANCEIQ_PRO_DOCUMENTATION.docx')
    print("✅ Word document created: FINANCEIQ_PRO_DOCUMENTATION.docx")

if __name__ == "__main__":
    try:
        create_word_document()
    except Exception as e:
        print(f"❌ Error creating Word document: {str(e)}")
        print("\nAlternative: Use Pandoc to convert:")
        print("pandoc FINANCEIQ_PRO_COMPREHENSIVE_DOCUMENTATION.md -o FINANCEIQ_PRO_DOCUMENTATION.docx")
